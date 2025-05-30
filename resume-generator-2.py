import os
import pdfplumber
import requests
from fpdf import FPDF
from datetime import datetime
import threading
import customtkinter as ctk
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
from docx import Document

# Set appearance
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")

# API and model config
TOGETHER_API_KEY = "8e0f28a685bf336d8ce153dd76ea065606cb7905edb0be1fbca5f2d84a031cbf"
LLAMA_MODEL = "mistralai/Mixtral-8x7B-Instruct-v0.1"
FONT_PATH = r"C:\Users\ElijahBabs\Downloads\Roboto (1)\Roboto-Italic-VariableFont_wdth,wght.ttf"


class ResumeTool:
    def __init__(self, pdf_path: str, job_description: str):
        self.pdf_path = pdf_path
        self.job_description = job_description
        self.resume_text = self.extract_text_from_pdf()


    def extract_resume_text(self) -> str:
        if self.resume_path.lower().endswith(".pdf"):
            return self.extract_text_from_pdf()
        elif self.resume_path.lower().endswith(".docx"):
            return self.extract_text_from_docx()
        

    def extract_text_from_pdf(self) -> str:
        text = ""
        with pdfplumber.open(self.pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    
    def extract_text_from_docx(self) -> str:
        doc = Document(self.resume_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def generate_with_llama(self, prompt: str) -> str:
        url = "https://api.together.xyz/inference"
        headers = {
            "Authorization": f"Bearer {TOGETHER_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": LLAMA_MODEL,
            "prompt": prompt,
            "max_tokens": 600,
            "temperature": 0.0,
            "top_p": 0.6
        }
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            return response.json()["choices"][0]["text"].strip()
        else:
            return f"Error: {response.status_code}\n{response.text}"

    def generate_tailored_resume(self) -> str:
        prompt = f"Tailor this resume to the job below...\nResume: {self.resume_text}\nJob: {self.job_description}"
        return self.clean_generated_resume(self.generate_with_llama(prompt))

    def generate_cover_letter(self) -> str:
        prompt = f"Write a short cover letter for the job below...\nResume: {self.resume_text}\nJob: {self.job_description}"
        return self.generate_with_llama(prompt)

    def clean_generated_resume(self, text: str) -> str:
        sections = ["PROFESSIONAL SUMMARY", "WORK EXPERIENCE", "SKILLS", "EDUCATION"]
        seen = set()
        clean_lines = []
        inside = False
        for line in text.split('\n'):
            stripped = line.strip()
            if any(sec in stripped.upper() for sec in sections):
                inside = True
                header = next((sec for sec in sections if sec in stripped.upper()), None)
                if header in seen:
                    break
                seen.add(header)
            if inside or stripped:
                clean_lines.append(stripped)
        return '\n'.join(clean_lines).strip()

    def prepare_paragraphs(self, text: str) -> list:
        return [line.strip() for line in text.split('\n') if line.strip()]

    def save_to_pdf(self, text: str, path: str):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=10)
        pdf.set_margins(left=15, top=10, right=15)
        pdf.add_font("Roboto", fname=FONT_PATH, uni=True)
        pdf.set_font("Roboto", size=11)
        width = pdf.w - pdf.l_margin - pdf.r_margin
        for para in self.prepare_paragraphs(text):
            pdf.multi_cell(width, 6.5, para)
            pdf.ln(3)
        pdf.output(path)

    def save_to_txt(self, text: str, path: str):
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)

    def save_to_docx(self, text: str, path: str):
        doc = Document()
        for para in self.prepare_paragraphs(text):
            doc.add_paragraph(para)
        doc.save(path)


class ResumeApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("AI Resume + Cover Letter Generator")
        self.geometry("950x750")
        self.resume_path = ""
        self.save_folder = ""

        self.create_widgets()

    def create_widgets(self):
        ctk.CTkLabel(self, text="Upload Resume (PDF):").grid(row=0, column=0, padx=10, pady=10, sticky="w")
        ctk.CTkButton(self, text="Browse...", command=self.browse_resume).grid(row=0, column=1, sticky="w")
        self.file_label = ctk.CTkLabel(self, text="No file selected", text_color="gray")
        self.file_label.grid(row=0, column=2, sticky="w")

        ctk.CTkLabel(self, text="Job Description:").grid(row=1, column=0, padx=10, pady=5, sticky="nw")
        self.job_text = ScrolledText(self, width=100, height=8)
        self.job_text.grid(row=2, column=0, columnspan=3, pady=5, padx=10)

        ctk.CTkButton(self, text="Choose Save Folder", command=self.choose_save_folder).grid(row=3, column=0, pady=5)
        ctk.CTkButton(self, text="Generate Resume", command=self.generate_resume).grid(row=3, column=1, pady=5)
        ctk.CTkButton(self, text="Generate Cover Letter", command=self.generate_cover_letter).grid(row=3, column=2, pady=5)

        ctk.CTkLabel(self, text="Resume Preview:").grid(row=4, column=0, sticky="w", padx=10)
        self.resume_preview = ScrolledText(self, width=100, height=10)
        self.resume_preview.grid(row=5, column=0, columnspan=3, padx=10, pady=5)

        ctk.CTkLabel(self, text="Cover Letter Preview:").grid(row=6, column=0, sticky="w", padx=10)
        self.cover_preview = ScrolledText(self, width=100, height=10)
        self.cover_preview.grid(row=7, column=0, columnspan=3, padx=10, pady=5)

        self.format_var = ctk.StringVar(value="PDF")
        self.format_menu = ctk.CTkOptionMenu(self, values=["PDF", "TXT", "DOCX"], variable=self.format_var)
        self.format_menu.grid(row=8, column=1, padx=10)

        ctk.CTkButton(self, text="Download Resume", command=self.save_edited_resume).grid(row=8, column=0, pady=10)
        ctk.CTkButton(self, text="Download Cover Letter", command=self.save_edited_cover_letter).grid(row=8, column=2, pady=10)

    def browse_resume(self):
        path = filedialog.askopenfilename(filetypes=[("Document Files", "*.pdf *.docx"), ("PDF Files", "*.pdf"), ("Word Documents", "*.docx")])
        if path:
            self.resume_path = path
            self.file_label.configure(text=os.path.basename(path), text_color="green")

    def choose_save_folder(self):
        folder = filedialog.askdirectory(title="Select Folder to Save Files")
        if folder:
            self.save_folder = folder

    def generate_resume(self):
        if not self.resume_path or not self.job_text.get("1.0", "end").strip():
            messagebox.showerror("Error", "Upload a resume and enter job description.")
            return

        def run():
            try:
                tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                result = tool.generate_tailored_resume()
                self.resume_preview.delete("1.0", "end")
                self.resume_preview.insert("end", result)
                messagebox.showinfo("Success", "Resume generated successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate resume:\n{e}")

        threading.Thread(target=run).start()

    def generate_cover_letter(self):
        if not self.resume_path or not self.job_text.get("1.0", "end").strip():
            messagebox.showerror("Error", "Upload a resume and enter job description.")
            return

        def run():
            try:
                tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                result = tool.generate_cover_letter()
                self.cover_preview.delete("1.0", "end")
                self.cover_preview.insert("end", result)
                messagebox.showinfo("Success", "Cover letter generated successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate cover letter:\n{e}")

        threading.Thread(target=run).start()

    def save_edited_resume(self):
        if not self.save_folder:
            self.choose_save_folder()
        if self.save_folder:
            text = self.resume_preview.get("1.0", "end").strip()
            if text:
                try:
                    tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    format_choice = self.format_var.get()
                    path = os.path.join(self.save_folder, f"edited_resume_{timestamp}.{format_choice.lower()}")
                    getattr(tool, f"save_to_{format_choice.lower()}")(text, path)
                    messagebox.showinfo("Success", f"Resume saved successfully:\n{path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save resume:\n{e}")

    def save_edited_cover_letter(self):
        if not self.save_folder:
            self.choose_save_folder()
        if self.save_folder:
            text = self.cover_preview.get("1.0", "end").strip()
            if text:
                try:
                    tool = ResumeTool(self.resume_path, self.job_text.get("1.0", "end").strip())
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    format_choice = self.format_var.get()
                    path = os.path.join(self.save_folder, f"edited_cover_letter_{timestamp}.{format_choice.lower()}")
                    getattr(tool, f"save_to_{format_choice.lower()}")(text, path)
                    messagebox.showinfo("Success", f"Cover letter saved successfully:\n{path}")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to save cover letter:\n{e}")


if __name__ == "__main__":
    app = ResumeApp()
    app.mainloop()
